# frontend/upload.py
import os
import io
import sqlite3
import hashlib
from datetime import datetime,UTC

import streamlit as st
from PyPDF2 import PdfReader
from docx import Document

# -----------------------
# Config
# -----------------------
DATA_DIR = "data"
UPLOAD_DIR = os.path.join(DATA_DIR, "uploads")
DB_PATH = os.path.join(DATA_DIR, "uploads.db")
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
ALLOWED_EXT = {".txt", ".pdf", ".docx"}

# Ensure folders exist
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(DATA_DIR, exist_ok=True)

# -----------------------
# Database utilities
# -----------------------
def init_db():
    conn = sqlite3.connect(DB_PATH, detect_types=sqlite3.PARSE_DECLTYPES)
    cur = conn.cursor()
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS books (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id TEXT,
            title TEXT,
            author TEXT,
            chapter TEXT,
            filename TEXT,
            filepath TEXT,
            filesize INTEGER,
            filehash TEXT,
            pages INTEGER,
            uploaded_at TIMESTAMP,
            status TEXT,
            summary_id TEXT
        )
        """
    )
    conn.commit()
    conn.close()

def insert_book(record):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute(
        """
        INSERT INTO books (user_id, title, author, chapter, filename, filepath, filesize, filehash, pages, uploaded_at, status)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            record.get("user_id"),
            record.get("title"),
            record.get("author"),
            record.get("chapter"),
            record["filename"],
            record["filepath"],
            record["filesize"],
            record["filehash"],
            record.get("pages"),
            record["uploaded_at"],
            record["status"],
        ),
    )
    conn.commit()
    book_id = cur.lastrowid
    conn.close()
    return book_id

def get_all_books():
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("SELECT id, title, author, chapter, filename, filesize, uploaded_at, status, summary_id FROM books ORDER BY uploaded_at DESC")
    rows = cur.fetchall()
    conn.close()
    return rows

def delete_book_db(book_id):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("SELECT filepath FROM books WHERE id = ?", (book_id,))
    r = cur.fetchone()
    if r and r[0] and os.path.exists(r[0]):
        try:
            os.remove(r[0])
        except Exception:
            pass
    cur.execute("DELETE FROM books WHERE id = ?", (book_id,))
    conn.commit()
    conn.close()

def update_book_status(book_id, status, summary_id=None):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    if summary_id:
        cur.execute("UPDATE books SET status = ?, summary_id = ? WHERE id = ?", (status, summary_id, book_id))
    else:
        cur.execute("UPDATE books SET status = ? WHERE id = ?", (status, book_id))
    conn.commit()
    conn.close()

# -----------------------
# Helpers
# -----------------------
def get_extension(filename):
    return os.path.splitext(filename)[1].lower()

def compute_hash(file_bytes: bytes) -> str:
    h = hashlib.sha256()
    h.update(file_bytes)
    return h.hexdigest()

def save_uploaded_file_to_disk(uploaded_file, dest_path):
    # uploaded_file: Streamlit UploadedFile
    with open(dest_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

def read_txt_preview(file_bytes, limit=500):
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        try:
            text = file_bytes.decode("latin-1")
        except Exception:
            return None
    preview = text[:limit]
    return preview

def pdf_num_pages(file_stream) -> int:
    try:
        reader = PdfReader(file_stream)
        return len(reader.pages)
    except Exception:
        return None

def docx_basic_info(file_stream):
    try:
        doc = Document(file_stream)
        paragraph_count = sum(1 for _ in doc.paragraphs)
        first_par = doc.paragraphs[0].text if paragraph_count > 0 else ""
        return {"paragraphs": paragraph_count, "first_par": first_par}
    except Exception:
        return None

# -----------------------
# UI
# -----------------------
def main():
    st.set_page_config(page_title="Upload Book for Summarization", layout="wide")
    init_db()

    st.title("Upload Book for Summarization")

    with st.expander("Instructions", expanded=True):
        st.markdown(
            """
            Supported formats: **.txt, .pdf, .docx**  
            Maximum file size: **10 MB**.  
            Upload will validate the file and store it. After upload you can click **Upload & Process** to trigger the backend summarization workflow.
            """
        )

    col1, col2 = st.columns([2, 1])
    with col1:
        uploaded_file = st.file_uploader(
            "Choose a book file",
            type=["txt", "pdf", "docx"],
            help="Maximum file size: 10MB"
        )
    with col2:
        st.write("Metadata (optional)")
        title_input = st.text_input("Book Title")
        author_input = st.text_input("Author")
        chapter_input = st.text_input("Chapter / Section")

    # Validation & Preview area
    if uploaded_file is not None:
        filename = uploaded_file.name
        ext = get_extension(filename)
        filesize = uploaded_file.size

        # Basic validations
        if ext not in ALLOWED_EXT:
            st.error(f"Unsupported file extension: {ext}. Allowed: {', '.join(ALLOWED_EXT)}")
            st.stop()

        if filesize == 0:
            st.error("Uploaded file is empty.")
            st.stop()

        if filesize > MAX_FILE_SIZE:
            st.error(f"File too large ({filesize} bytes). Maximum is {MAX_FILE_SIZE} bytes.")
            st.stop()

        # Read bytes for validation + duplicate detection
        file_bytes = uploaded_file.read()
        filehash = compute_hash(file_bytes)
        # Reset buffer pointer for downstream readers
        uploaded_file.seek(0)

        # Check duplicate by filehash in DB
        conn = sqlite3.connect(DB_PATH)
        cur = conn.cursor()
        cur.execute("SELECT id, title, filename, uploaded_at FROM books WHERE filehash = ?", (filehash,))
        duplicate = cur.fetchone()
        conn.close()
        if duplicate:
            st.warning(f"A file with the same content was uploaded earlier ({duplicate[2]} on {duplicate[3]}).")
            if not st.checkbox("Proceed anyway and upload duplicate"):
                st.stop()

        st.subheader("File preview & info")
        st.write(f"**Filename:** {filename}")
        st.write(f"**Size:** {filesize} bytes")
        st.write(f"**Uploaded (in this form):** {datetime.now(UTC).isoformat()}")

        preview_cols = st.columns([1, 1])
        with preview_cols[0]:
            if ext == ".txt":
                preview = read_txt_preview(file_bytes, limit=500)
                if preview is None:
                    st.error("Could not decode TXT file. It may use an unsupported encoding or be corrupted.")
                else:
                    st.text_area("TXT preview (first 500 chars)", preview, height=200)
            elif ext == ".pdf":
                try:
                    # For PyPDF2 we need a bytes stream
                    stream = io.BytesIO(file_bytes)
                    pages = pdf_num_pages(stream)
                    if pages is None:
                        st.error("Could not read PDF. It might be corrupted or scanned (image-only).")
                    else:
                        st.write(f"Number of pages: **{pages}**")
                except Exception as e:
                    st.error("Error reading PDF: " + str(e))
            elif ext == ".docx":
                try:
                    stream = io.BytesIO(file_bytes)
                    info = docx_basic_info(stream)
                    if info is None:
                        st.error("Could not read DOCX. It might be corrupted.")
                    else:
                        st.write(f"Paragraphs: **{info['paragraphs']}**")
                        if info['first_par']:
                            st.text_area("First paragraph", info['first_par'], height=150)
                except Exception as e:
                    st.error("Error reading DOCX: " + str(e))

        # Pre-fill title from filename if title_input is empty
        if not title_input:
            inferred_title = os.path.splitext(filename)[0]
            st.info(f"Auto-filled title from filename: {inferred_title}")
            title_input = st.text_input("Book Title (edit if needed)", value=inferred_title)

        # Upload & Process button
        if st.button("Upload & Process"):
            # Save file to disk
            timestamp = datetime.now(UTC).isoformat()
            safe_filename = f"{int(datetime.now(UTC).timestamp())}_{filename}"
            dest_path = os.path.join(UPLOAD_DIR, safe_filename)

            try:
                with st.spinner("Saving file..."):
                    # write bytes to disk
                    with open(dest_path, "wb") as f:
                        f.write(file_bytes)

                # determine pages for storing metadata
                pages = None
                if ext == ".pdf":
                    try:
                        pages = pdf_num_pages(io.BytesIO(file_bytes))
                    except Exception:
                        pages = None
                elif ext == ".docx":
                    try:
                        info = docx_basic_info(io.BytesIO(file_bytes))
                        pages = info["paragraphs"] if info else None
                    except Exception:
                        pages = None

                rec = {
                    "user_id": None,
                    "title": title_input or inferred_title,
                    "author": author_input,
                    "chapter": chapter_input,
                    "filename": filename,
                    "filepath": dest_path,
                    "filesize": filesize,
                    "filehash": filehash,
                    "pages": pages,
                    "uploaded_at": timestamp,
                    "status": "uploaded"
                }
                book_id = insert_book(rec)
                st.success(f"File saved and metadata stored (book id: {book_id}).")

                # Show progress for file-handling step
                progress_bar = st.progress(0)
                for i in range(5):
                    progress_bar.progress((i + 1) * 20)
                progress_bar.empty()

                st.info("To start summarization, the frontend should call the backend orchestration endpoint for this book_id.")
                st.write("Example (placeholder):")
                st.code(f'POST /api/summarize?book_id={book_id}  # call your backend job runner/orchestrator', language="bash")

                # Optionally update status to processing here if you trigger backend immediately
                # update_book_status(book_id, "processing")
            except Exception as e:
                st.error(f"Failed to save file: {e}")

    # -----------------------
    # Upload history section
    # -----------------------
    st.markdown("---")
    st.header("Upload history / status")

    books = get_all_books()
    if not books:
        st.info("No uploads yet. Use the form above to add a book.")
    else:
        # Render a simple table and actions
        for row in books:
            book_id, title, author, chapter, filename, filesize, uploaded_at, status, summary_id = row
            with st.container():
                cols = st.columns([3, 1, 1, 1])
                cols[0].markdown(f"**{title}**  \n*{filename}*  \nAuthor: {author or '—'}  \nChapter: {chapter or '—'}")
                cols[1].markdown(f"Uploaded: {uploaded_at}")
                cols[2].markdown(f"Status: **{status}**")
                action_col = cols[3]
                # View summary button (enabled only if summary exists)
                if summary_id:
                    if action_col.button("View Summary", key=f"view_{book_id}"):
                        # Implement navigation to summary page or open a new tab with summary
                        # Placeholder: instruct user; in your app you'd call st.experimental_set_query_params or navigate to page
                        st.session_state.setdefault("navigate_to_summary", None)
                        st.session_state["navigate_to_summary"] = summary_id
                        st.success(f"Would navigate to summary {summary_id}. Implement page: frontend/summary.py")
                else:
                    action_col.write("")  # keep layout

                # Delete button
                if action_col.button("Delete", key=f"del_{book_id}"):
                    delete_book_db(book_id)
                    st.experimental_rerun()

        st.write("Tip: If you implement async background processing in backend, update the 'status' field using update_book_status(book_id, 'processing'/'completed'/'failed').")

    # -----------------------
    # End
    # -----------------------
